"""
Document Export Service
Handles exporting documents to various formats (PDF, DOCX, MD) with proper markdown formatting.
"""

import io
import markdown
from typing import BinaryIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
try:
    from weasyprint import HTML, CSS
    from weasyprint.text.fonts import FontConfiguration
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    HTML = None
    CSS = None
    FontConfiguration = None
import re


def export_to_markdown(content: str, title: str) -> bytes:
    """
    Export document content as Markdown.
    Simple passthrough since content is already in markdown format.
    """
    output = f"# {title}\n\n{content}"
    return output.encode('utf-8')


def export_to_pdf(content: str, title: str) -> bytes:
    """
    Export document to PDF with properly formatted markdown.
    Uses WeasyPrint to convert HTML (from markdown) to PDF.
    """
    # Convert markdown to HTML with extensions
    html_content = markdown.markdown(
        content,
        extensions=[
            'extra',  # Tables, fenced code blocks, etc.
            'codehilite',  # Syntax highlighting
            'nl2br',  # Newline to <br>
            'sane_lists',  # Better list handling
            'toc',  # Table of contents
        ]
    )

    # Create complete HTML document with styling
    html_template = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <meta charset="utf-8">
        <title>{title}</title>
    </head>
    <body>
        <h1 class="document-title">{title}</h1>
        <div class="content">
            {html_content}
        </div>
    </body>
    </html>
    """

    # CSS styling for better PDF appearance
    css_content = """
    @page {
        size: A4;
        margin: 2cm;
    }

    body {
        font-family: 'DejaVu Sans', Arial, sans-serif;
        font-size: 11pt;
        line-height: 1.6;
        color: #333;
    }

    .document-title {
        font-size: 24pt;
        font-weight: bold;
        margin-bottom: 1em;
        padding-bottom: 0.5em;
        border-bottom: 2px solid #333;
        color: #000;
    }

    h1 {
        font-size: 20pt;
        margin-top: 1.5em;
        margin-bottom: 0.5em;
        color: #000;
    }

    h2 {
        font-size: 16pt;
        margin-top: 1.2em;
        margin-bottom: 0.4em;
        color: #222;
    }

    h3 {
        font-size: 14pt;
        margin-top: 1em;
        margin-bottom: 0.3em;
        color: #333;
    }

    h4, h5, h6 {
        font-size: 12pt;
        margin-top: 0.8em;
        margin-bottom: 0.3em;
        color: #444;
    }

    p {
        margin: 0.5em 0;
        text-align: justify;
    }

    code {
        background-color: #f4f4f4;
        padding: 2px 6px;
        border-radius: 3px;
        font-family: 'DejaVu Sans Mono', 'Courier New', monospace;
        font-size: 9pt;
    }

    pre {
        background-color: #f8f8f8;
        border: 1px solid #ddd;
        border-left: 3px solid #4CAF50;
        padding: 1em;
        overflow-x: auto;
        border-radius: 4px;
        margin: 1em 0;
    }

    pre code {
        background-color: transparent;
        padding: 0;
        font-size: 9pt;
    }

    blockquote {
        border-left: 4px solid #2196F3;
        padding-left: 1em;
        margin: 1em 0;
        color: #555;
        font-style: italic;
        background-color: #f9f9f9;
        padding: 0.5em 1em;
    }

    ul, ol {
        margin: 0.5em 0;
        padding-left: 2em;
    }

    li {
        margin: 0.3em 0;
    }

    table {
        border-collapse: collapse;
        width: 100%;
        margin: 1em 0;
    }

    th, td {
        border: 1px solid #ddd;
        padding: 8px 12px;
        text-align: left;
    }

    th {
        background-color: #4CAF50;
        color: white;
        font-weight: bold;
    }

    tr:nth-child(even) {
        background-color: #f9f9f9;
    }

    a {
        color: #2196F3;
        text-decoration: none;
    }

    a:hover {
        text-decoration: underline;
    }

    hr {
        border: none;
        border-top: 2px solid #ddd;
        margin: 2em 0;
    }

    img {
        max-width: 100%;
        height: auto;
        display: block;
        margin: 1em 0;
    }

    .codehilite {
        background-color: #f8f8f8;
        border-radius: 4px;
        padding: 1em;
        margin: 1em 0;
    }
    """

    if not WEASYPRINT_AVAILABLE:
        raise NotImplementedError("PDF export is not available. Missing system dependencies (weasyprint/cairo).")

    # Generate PDF
    font_config = FontConfiguration()
    pdf_bytes = HTML(string=html_template).write_pdf(
        stylesheets=[CSS(string=css_content)],
        font_config=font_config
    )

    return pdf_bytes


def export_to_docx(content: str, title: str) -> bytes:
    """
    Export document to DOCX with properly formatted markdown.
    Parses markdown and creates styled Word document.
    """
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Add title
    title_paragraph = doc.add_heading(title, level=0)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Parse markdown content line by line
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_block_lines = []
    in_list = False
    list_level = 0

    while i < len(lines):
        line = lines[i]

        # Handle code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_block_lines)
                p = doc.add_paragraph(code_text)
                p.style = 'Quote'
                # Set monospace font
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_block_lines = []
                in_code_block = False
            else:
                # Start code block
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Handle headings
        if line.startswith('#'):
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2)
                doc.add_heading(heading_text, level=level)
                i += 1
                continue

        # Handle horizontal rules
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 50)
            i += 1
            continue

        # Handle blockquotes
        if line.strip().startswith('>'):
            quote_text = line.strip()[1:].strip()
            p = doc.add_paragraph(quote_text, style='Quote')
            i += 1
            continue

        # Handle unordered lists
        if re.match(r'^\s*[-*+]\s+', line):
            list_match = re.match(r'^(\s*)([-*+])\s+(.+)$', line)
            if list_match:
                indent = len(list_match.group(1))
                text = list_match.group(3)
                p = doc.add_paragraph(text, style='List Bullet')
                i += 1
                continue

        # Handle ordered lists
        if re.match(r'^\s*\d+\.\s+', line):
            list_match = re.match(r'^(\s*)(\d+)\.\s+(.+)$', line)
            if list_match:
                text = list_match.group(3)
                p = doc.add_paragraph(text, style='List Number')
                i += 1
                continue

        # Handle tables (basic support)
        if '|' in line and line.strip().startswith('|'):
            # Collect table rows
            table_rows = []
            while i < len(lines) and '|' in lines[i]:
                row_line = lines[i].strip()
                if not re.match(r'^\|[\s\-:]+\|$', row_line):  # Skip separator row
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    table_rows.append(cells)
                i += 1

            if table_rows:
                # Create table
                table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
                table.style = 'Light Grid Accent 1'

                for row_idx, row_data in enumerate(table_rows):
                    for col_idx, cell_data in enumerate(row_data):
                        cell = table.rows[row_idx].cells[col_idx]
                        cell.text = cell_data
                        # Bold first row (header)
                        if row_idx == 0:
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.bold = True
            continue

        # Handle regular paragraphs with inline formatting
        if line.strip():
            p = doc.add_paragraph()

            # Parse inline markdown
            text = line

            # Split by various markdown patterns
            # This is a simplified parser - a full parser would be more complex
            parts = [text]

            # Handle bold (**text** or __text__)
            for part in parts[:]:
                if '**' in part or '__' in part:
                    segments = re.split(r'(\*\*.*?\*\*|__.*?__)', part)
                    parts.remove(part)
                    parts.extend(segments)

            # Handle italic (*text* or _text_)
            new_parts = []
            for part in parts:
                if re.search(r'(?<!\*)\*(?!\*).*?(?<!\*)\*(?!\*)', part) or re.search(r'(?<!_)_(?!_).*?(?<!_)_(?!_)', part):
                    segments = re.split(r'((?<!\*)\*(?!\*).*?(?<!\*)\*(?!\*)|(?<!_)_(?!_).*?(?<!_)_(?!_))', part)
                    new_parts.extend(segments)
                else:
                    new_parts.append(part)
            parts = new_parts

            # Handle inline code (`code`)
            new_parts = []
            for part in parts:
                if '`' in part:
                    segments = re.split(r'(`[^`]+?`)', part)
                    new_parts.extend(segments)
                else:
                    new_parts.append(part)
            parts = new_parts

            # Add runs with appropriate formatting
            for part in parts:
                if not part:
                    continue

                run = p.add_run(part)

                # Apply formatting based on markdown syntax
                if part.startswith('**') and part.endswith('**'):
                    run.text = part[2:-2]
                    run.font.bold = True
                elif part.startswith('__') and part.endswith('__'):
                    run.text = part[2:-2]
                    run.font.bold = True
                elif re.match(r'^\*[^*]+\*$', part):
                    run.text = part[1:-1]
                    run.font.italic = True
                elif re.match(r'^_[^_]+_$', part):
                    run.text = part[1:-1]
                    run.font.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    run.text = part[1:-1]
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(220, 50, 47)
        else:
            # Empty line
            doc.add_paragraph()

        i += 1

    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)

    return docx_bytes.read()
